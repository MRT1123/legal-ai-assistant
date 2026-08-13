"""
文档导出服务
支持将生成的法律文书导出为 Word (.docx) 或 PDF (.pdf) 格式
"""

import os
from fpdf import FPDF


# 导出文件存放目录
EXPORT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "data", "exports")


def _ensure_export_dir():
    """确保导出目录存在"""
    os.makedirs(EXPORT_DIR, exist_ok=True)


def _clean_markdown(text: str) -> str:
    """
    清理 Markdown 标记，转为纯文本
    PDF 导出时需要去掉 Markdown 语法
    """
    lines = text.split('\n')
    cleaned = []
    for line in lines:
        # 去掉标题标记（# ## ###）
        if line.startswith('### '):
            line = line[4:]
        elif line.startswith('## '):
            line = line[3:]
        elif line.startswith('# '):
            line = line[2:]
        # 去掉加粗标记 **text**
        line = line.replace('**', '')
        # 去掉斜体标记 *text*
        import re
        line = re.sub(r'\*(.+?)\*', r'\1', line)
        cleaned.append(line)
    return '\n'.join(cleaned)


def export_to_word(content: str, title: str = "法律文书", filename: str = None) -> str:
    """
    导出为 Word 文档 (.docx)
    
    参数：
        content:   文档正文内容（支持纯文本）
        title:     文档标题
        filename:  自定义文件名（不含扩展名），不传则自动生成
    
    返回：
        生成文件的完整路径
    """
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    _ensure_export_dir()

    if not filename:
        import uuid
        filename = f"文书_{uuid.uuid4().hex[:8]}"

    filepath = os.path.join(EXPORT_DIR, f"{filename}.docx")

    doc = Document()

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # 添加标题
    heading = doc.add_heading(level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(title)
    run.font.size = Pt(22)
    run.font.color.rgb = None  # 使用默认黑色

    # 添加正文内容（按段落写入）
    paragraphs = content.split('\n\n')
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue

        # 处理 Markdown 标题
        if para_text.startswith('# '):
            h = doc.add_heading(level=2)
            h.add_run(para_text[2:].strip())
        elif para_text.startswith('## '):
            h = doc.add_heading(level=2)
            h.add_run(para_text[3:].strip())
        elif para_text.startswith('### '):
            h = doc.add_heading(level=3)
            h.add_run(para_text[4:].strip())
        else:
            # 普通段落，处理列表项
            for line in para_text.split('\n'):
                line = line.strip()
                if not line:
                    continue
                # 去掉加粗标记
                line = line.replace('**', '')
                doc.add_paragraph(line)

    doc.save(filepath)
    return filepath


def export_to_pdf(content: str, title: str = "法律文书", filename: str = None) -> str:
    """
    导出为 PDF 文档
    使用 Windows 自带的微软雅黑字体（msyh.ttc）渲染中文
    
    参数：
        content:   文档正文内容（支持纯文本）
        title:     文档标题
        filename:  自定义文件名（不含扩展名），不传则自动生成
    
    返回：
        生成文件的完整路径
    """
    _ensure_export_dir()

    if not filename:
        import uuid
        filename = f"文书_{uuid.uuid4().hex[:8]}"

    filepath = os.path.join(EXPORT_DIR, f"{filename}.pdf")

    # 查找 Windows 中文字体
    font_path = _find_chinese_font()
    if not font_path:
        raise RuntimeError(
            "未找到中文字体文件。请确保系统中存在微软雅黑字体（msyh.ttc 或 simsun.ttc）。"
        )

    # 创建 PDF
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # 注册中文字体
    pdf.add_font("Chinese", "", font_path, uni=True)
    pdf.add_font("Chinese", "B", font_path, uni=True)

    # 绘制标题（居中，大号字）
    pdf.set_font("Chinese", "B", 20)
    pdf.cell(0, 15, title, ln=True, align="C")
    pdf.ln(10)

    # 绘制正文
    pdf.set_font("Chinese", "", 11)
    cleaned_text = _clean_markdown(content)

    for line in cleaned_text.split('\n'):
        line = line.strip()
        if not line:
            pdf.ln(4)  # 空行
            continue
        # multi_cell 支持自动换行
        pdf.multi_cell(0, 7, line)

    pdf.output(filepath)
    return filepath


def _find_chinese_font() -> str:
    """
    查找系统中可用的中文字体文件
    优先顺序：微软雅黑 → 宋体 → 黑体
    """
    # Windows 字体目录
    fonts_dir = os.path.join(os.environ.get("WINDIR", "C:\\Windows"), "Fonts")

    # 按优先级查找
    candidates = [
        "msyh.ttc",      # 微软雅黑
        "msyhbd.ttc",    # 微软雅黑粗体
        "simsun.ttc",    # 宋体
        "simhei.ttf",    # 黑体
        "simfang.ttf",   # 仿宋
    ]

    for font_name in candidates:
        font_path = os.path.join(fonts_dir, font_name)
        if os.path.exists(font_path):
            return font_path

    return None
